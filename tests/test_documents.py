from __future__ import annotations

import io

from docx import Document

from llamatui.documents import DocumentResult, extract_document


def _docx_bytes(build) -> bytes:
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_result_constructors_set_status_and_payload():
    assert DocumentResult.extracted("hi").status == "extracted"
    assert DocumentResult.extracted("hi").text == "hi"
    assert DocumentResult.not_a_document().status == "not_a_document"
    assert DocumentResult.needs_ocr("scan").status == "needs_ocr"
    assert DocumentResult.needs_ocr("scan").reason == "scan"
    assert DocumentResult.failed("boom").status == "failed"
    assert DocumentResult.failed("boom").reason == "boom"


def test_unknown_extension_is_not_a_document():
    assert extract_document(b"hello world", "notes.txt").status == "not_a_document"


def test_docx_with_bad_magic_is_not_a_document():
    # .docx extension but not a ZIP/PK container
    assert extract_document(b"not a zip", "fake.docx").status == "not_a_document"


def test_docx_extracts_paragraphs_and_tables_in_order():
    def build(doc):
        doc.add_paragraph("Intro line")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A"
        table.cell(0, 1).text = "B"
        table.cell(1, 0).text = "C"
        table.cell(1, 1).text = "D"
        doc.add_paragraph("Outro line")

    result = extract_document(_docx_bytes(build), "doc.docx")
    assert result.status == "extracted"
    lines = result.text.splitlines()
    assert lines[0] == "Intro line"
    assert "A | B" in lines
    assert "C | D" in lines
    assert lines[-1] == "Outro line"
    # table appears between the two paragraphs, in document order
    assert lines.index("A | B") > lines.index("Intro line")
    assert lines.index("A | B") < lines.index("Outro line")


def test_empty_docx_fails_with_reason():
    result = extract_document(_docx_bytes(lambda doc: None), "empty.docx")
    assert result.status == "failed"
    assert result.reason


def test_missing_python_docx_reports_failed(monkeypatch):
    import llamatui.documents as documents

    monkeypatch.setattr(documents, "docx", None)
    result = extract_document(_docx_bytes(lambda doc: doc.add_paragraph("x")), "doc.docx")
    assert result.status == "failed"
    assert "python-docx" in result.reason


# ---------------------------------------------------------------------------
# PDF tests
# ---------------------------------------------------------------------------


def _pdf_with_text(text: str) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("helvetica", size=14)
    pdf.cell(text=text)
    return bytes(pdf.output())


def _pdf_blank_page() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _pdf_encrypted() -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt("secret")  # non-empty password -> not readable with ""
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def test_pdf_with_bad_magic_is_not_a_document():
    assert extract_document(b"not a pdf", "fake.pdf").status == "not_a_document"


def test_pdf_extracts_text_layer():
    result = extract_document(_pdf_with_text("Hello from PDF"), "doc.pdf")
    assert result.status == "extracted"
    assert "Hello from PDF" in result.text


def test_image_only_pdf_needs_ocr():
    result = extract_document(_pdf_blank_page(), "scan.pdf")
    assert result.status == "needs_ocr"
    assert result.reason


def test_encrypted_pdf_fails():
    result = extract_document(_pdf_encrypted(), "locked.pdf")
    assert result.status == "failed"
    assert "encrypted" in result.reason.lower()


def test_missing_pypdf_reports_failed(monkeypatch):
    import llamatui.documents as documents

    monkeypatch.setattr(documents, "pypdf", None)
    result = extract_document(_pdf_with_text("x"), "doc.pdf")
    assert result.status == "failed"
    assert "pypdf" in result.reason
